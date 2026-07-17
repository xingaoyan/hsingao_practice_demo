#!/usr/bin/env python3
"""
DOCX 解析器
"""

from pathlib import Path
from typing import Dict, List, Optional, Any
from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph

from .base import BaseParser, ParseResult, PageInfo


class DOCXParser(BaseParser):
    """DOCX 解析器"""
    
    def parse(self, file_path: Path, document_id: str) -> ParseResult:
        """解析 DOCX 文件"""
        self.clear_warnings()
        
        try:
            doc = Document(str(file_path))
        except Exception as e:
            self.add_warning(f"无法打开 DOCX 文件: {e}")
            return self._create_empty_result(file_path, document_id)
        
        pages = []
        metadata = self._extract_metadata(doc, file_path)
        
        # 提取内容
        content = self._extract_content(doc)
        
        # 将内容分页（DOCX 没有明确的页面概念，按段落分组）
        page_size = 50  # 每页大约50个段落
        for i in range(0, len(content), page_size):
            page_content = content[i:i + page_size]
            page_text = "\n".join(page_content)
            
            pages.append(PageInfo(
                page_number=i // page_size + 1,
                text=page_text,
                images=[],
                tables=[]
            ))
        
        # 提取标题
        title = self._extract_title(doc, file_path)
        
        return ParseResult(
            document_id=document_id,
            title=title,
            source_path=str(file_path),
            file_type="docx",
            pages=pages,
            metadata=metadata,
            warnings=self.warnings
        )
    
    def _extract_content(self, doc: DocumentType) -> List[str]:
        """提取文档内容"""
        content = []
        
        for element in doc.element.body:
            if element.tag.endswith('}p'):  # 段落
                paragraph = Paragraph(element, doc)
                text = paragraph.text.strip()
                if text:
                    # 保留标题样式信息
                    style_name = paragraph.style.name if paragraph.style else ""
                    if "Heading" in style_name:
                        level = style_name.replace("Heading ", "")
                        content.append(f"{'#' * int(level)} {text}")
                    else:
                        content.append(text)
            
            elif element.tag.endswith('}tbl'):  # 表格
                table = Table(element, doc)
                table_text = self._extract_table_text(table)
                if table_text:
                    content.append(table_text)
        
        return content
    
    def _extract_table_text(self, table: Table) -> str:
        """提取表格文本"""
        try:
            rows = []
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cells.append(cell.text.strip())
                rows.append(cells)
            
            if not rows:
                return ""
            
            # 创建 Markdown 表格
            table_lines = []
            table_lines.append("| " + " | ".join(rows[0]) + " |")
            table_lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
            for row in rows[1:]:
                table_lines.append("| " + " | ".join(row) + " |")
            
            return "\n".join(table_lines)
            
        except Exception as e:
            self.add_warning(f"提取表格时出错: {e}")
            return ""
    
    def _extract_metadata(self, doc: DocumentType, file_path: Path) -> Dict[str, Any]:
        """提取文档元数据"""
        metadata = self.get_metadata(file_path)
        
        try:
            # 提取核心属性
            core_properties = doc.core_properties
            if core_properties:
                metadata.update({
                    "title": core_properties.title or "",
                    "author": core_properties.author or "",
                    "subject": core_properties.subject or "",
                    "keywords": core_properties.keywords or "",
                    "comments": core_properties.comments or "",
                    "category": core_properties.category or "",
                    "created": str(core_properties.created) if core_properties.created else "",
                    "modified": str(core_properties.modified) if core_properties.modified else "",
                    "last_modified_by": core_properties.last_modified_by or ""
                })
        except Exception as e:
            self.add_warning(f"提取元数据时出错: {e}")
        
        return metadata
    
    def _extract_title(self, doc: DocumentType, file_path: Path) -> str:
        """提取文档标题"""
        # 首先尝试从元数据获取
        try:
            if doc.core_properties.title:
                return doc.core_properties.title
        except:
            pass
        
        # 尝试从第一个段落提取
        try:
            if doc.paragraphs:
                first_paragraph = doc.paragraphs[0]
                if first_paragraph.text.strip():
                    return first_paragraph.text.strip()[:100]
        except:
            pass
        
        # 使用文件名作为标题
        return file_path.stem
    
    def _create_empty_result(self, file_path: Path, document_id: str) -> ParseResult:
        """创建空结果"""
        return ParseResult(
            document_id=document_id,
            title=file_path.stem,
            source_path=str(file_path),
            file_type="docx",
            pages=[],
            metadata=self.get_metadata(file_path),
            warnings=self.warnings
        )